# Nessus results parser by Ivan Morales
# and python 3.x
import argparse
import glob
import os
import sys

from docx import Document
from lxml import etree
from collections import OrderedDict


banner="""
  _   _                           ____                                           
 | \ | | ___  ___ ___ _   _ ___  |  _ \ __ _ _ __ ___  ___ _ __                  
 |  \| |/ _ \/ __/ __| | | / __| | |_) / _` | '__/ __|/ _ \ '__|                 
 | |\  |  __/\__ \__ \ |_| \__ \ |  __/ (_| | |  \__ \  __/ |                    
 |_|_\_|\___||___/___/\__,_|___/ |_|___\__,_|_|  |___/\___|_|      _             
 |  _ \ ___ _ __   ___  _ __| |_   / ___| ___ _ __   ___ _ __ __ _| |_ ___  _ __ 
 | |_) / _ \ '_ \ / _ \| '__| __| | |  _ / _ \ '_ \ / _ \ '__/ _` | __/ _ \| '__|
 |  _ <  __/ |_) | (_) | |  | |_  | |_| |  __/ | | |  __/ | | (_| | || (_) | |   
 |_| \_\___| .__/ \___/|_|   \__|  \____|\___|_| |_|\___|_|  \__,_|\__\___/|_|   
           |_|                                                                   
        ==========================================
          - Release date: 2021-08-13
          - Script Parser nessus file to wordx
          - Version 1.0.0
        ===========================================

                        Written by:
                        Ivan Morales
                    https://github.com/ivenmori/ 
"""


class NessusParser:
    def __init__( self, severity):
        self.severity = severity

    def xmlreader(self,file):

        f= open(file, 'r')
        xml_content = f.read()
        f.close()
        return xml_content

    def get_vulners_from_xmls(self,nessus_list):
        vulnerabilities = dict()
        vulns_sort = dict()
        single_params = ["agent", "cvss3_base_score", "cvss3_temporal_score", "cvss3_temporal_vector", "cvss3_vector",
                        "cvss_base_score", "cvss_temporal_score", "cvss_temporal_vector", "cvss_vector", "description",
                        "exploit_available", "exploitability_ease", "exploited_by_nessus", "fname", "in_the_news",
                        "patch_publication_date", "plugin_modification_date", "plugin_name", "plugin_publication_date",
                        "plugin_type", "script_version", "see_also", "solution", "synopsis", "vuln_publication_date",
                        "compliance",
                        "{http://www.nessus.org/cm}compliance-check-id",
                        "{http://www.nessus.org/cm}compliance-check-name",
                        "{http://www.nessus.org/cm}audit-file",
                        "{http://www.nessus.org/cm}compliance-info",
                        "{http://www.nessus.org/cm}compliance-result",
                        "{http://www.nessus.org/cm}compliance-see-also"]
        for nessus_file in nessus_list:
            print ('Parsing nessus file: ' + nessus_file)
            xml_content=self.xmlreader(nessus_file)

            p = etree.XMLParser(huge_tree=True)
            root = etree.fromstring(text=xml_content, parser=p)
            for block in root:
                if block.tag == "Report":
                    for report_host in block:
                        host_properties_dict = dict()
                        for report_item in report_host:
                            if report_item.tag == "HostProperties":
                                for host_properties in report_item:
                                    host_properties_dict[host_properties.attrib['name']] = host_properties.text
        ### This IF is to run the tree one time.                            
                            if report_item.tag == "ReportItem": 
                                if 'severity' in report_item.attrib:
                                    if int(report_item.attrib['severity']) >= self.severity:
                                        vulner_struct = dict()
                                        vulner_struct['cvss_sort']=0.0
                                        vulner_struct['port'] = report_item.attrib['port']
                                        vulner_struct['pluginName'] = report_item.attrib['pluginName']
                                        vulner_id = report_item.attrib['pluginName']
                                        vulner_struct['pluginFamily'] = report_item.attrib['pluginFamily']
                                        vulner_struct['pluginID'] = report_item.attrib['pluginID']
                                        vulner_struct['svc_name'] = report_item.attrib['svc_name']
                                        vulner_struct['protocol'] = report_item.attrib['protocol']


                                        for param in report_item:
                                            if param.tag == "risk_factor":
                                                risk_factor = param.text
                                                vulner_struct['host'] = report_host.attrib['name']
                                                vulner_struct['risk_factor'] = risk_factor
                                            elif param.tag == "plugin_output":
                                                if not "plugin_output" in vulner_struct:
                                                    vulner_struct["plugin_output"] = list()
                                                if not param.text in vulner_struct["plugin_output"]:
                                                    vulner_struct["plugin_output"].append(param.text)                                        
                                            else:
                                                if not param.tag in single_params:
                                                    if not param.tag in vulner_struct:
                                                        vulner_struct[param.tag] = list()
                                                    if not isinstance(vulner_struct[param.tag], list):
                                                        vulner_struct[param.tag] = [vulner_struct[param.tag]]
                                                    if not param.text in vulner_struct[param.tag]:
                                                        vulner_struct[param.tag].append(param.text)
                                                else:
                                                    vulner_struct[param.tag] = param.text

                                        for param in host_properties_dict:
                                            vulner_struct[param] = host_properties_dict[param]
                                        protocol = dict()
                                        protocol[vulner_struct['protocol']]=vulner_struct['port']
                                        host_ports = dict()
                                        host_ports[vulner_struct['host']] =protocol
                                        vulner_struct['affected']  = host_ports
                                        # Use only the CVSS 2 or CVSS 3 with values 0.0 to 10.0
                                        cvss2=0.0
                                        cvss3=0.0
                                        if ('cvss_base_score' in vulner_struct) or ('cvss3_base_score'in vulner_struct):
                                            if vulner_struct.get('cvss_base_score') is not None:
                                                cvss2 = float(vulner_struct['cvss_base_score']) 
                                            if vulner_struct.get('cvss3_base_score') is not None:
                                                cvss3 = float(vulner_struct['cvss3_base_score'])
                                            if (cvss2 < cvss3) :
                                                vulner_struct['cvss_sort'] = cvss3
                                            else :
                                                vulner_struct['cvss_sort'] = cvss2 
                                            
                                            if not vulner_id in vulnerabilities:
                                                vulnerabilities[vulner_id] = vulner_struct

                                            else:
                                                if not vulner_struct['host'] in vulnerabilities[vulner_id]['affected']:
                                                    vulnerabilities[vulner_id]['affected'][vulner_struct['host']]= protocol
                                                else:
                                                    if(not vulner_struct['protocol'] in vulnerabilities[vulner_id]['affected'][vulner_struct['host']]):
                                                        vulnerabilities[vulner_id]['affected'][vulner_struct['host']][vulner_struct['protocol']]=vulner_struct['port']
                                                    else:
                                                        if not vulner_struct['port'] in vulnerabilities[vulner_id]['affected'][vulner_struct['host']][vulner_struct['protocol']]:
                                                            vulnerabilities[vulner_id]['affected'][vulner_struct['host']][vulner_struct['protocol']]= vulner_struct['port'] +','+ vulnerabilities[vulner_id]['affected'][vulner_struct['host']][vulner_struct['protocol']]
        ##sorted(vulnerabilities.items() , key=lambda vi : vi[1]["cvss_sort"])
        ##sorted(vulnerabilities..items(), key=lambda x: x[1]['severity'], reverse=False)
        return(OrderedDict(sorted(vulnerabilities.items(), key=lambda x: x[1]['cvss_sort'], reverse=True)))


class Reporter:
    def __init__( self, ofile):
        self.ofile = ofile
        self.count = 0
        self.document = Document()

    def vulnreader(self,vulnerabilities):
        for vulner_id in vulnerabilities:
            if ('cvss_base_score' in vulners[vulner_id]) or ('cvss3_base_score'in vulners[vulner_id]):
                self.savedocx(vulners[vulner_id])
        self.document.save(self.ofile)
        print("[+] Created word file " + self.ofile)           

    def savedocx(self, vuln):
        self.document.add_paragraph('')
        
        self.document.add_heading(vuln['pluginName'], level=2)

        #Default values in case they have no values assigned in the xml
        cvss3 = ""
        cvss2 = ""    
        see_also = ""
        see_also = ""
        plugin_output = ""
        
        if vuln.get('cvss_base_score') is not None:
            cvss2 = vuln['cvss_base_score'] + " " + vuln['cvss_vector']

        if vuln.get('cvss3_base_score') is not None:
            cvss3 = vuln['cvss3_base_score'] + " " + vuln['cvss3_vector']

        if vuln.get('see_also') is not None:
            see_also = vuln['see_also']

        if vuln.get('plugin_output') is not None:
            plugin_output =vuln.get('plugin_output')
        
        affected =""
        for k,v in vuln['affected'].items():
            affected = affected + k + " on port "
            for  ki, vi in v.items():
                affected = affected + vi+'/'+ki
            affected = affected + '\n'
        
    

        records = (
###         ('Name:',vuln['pluginName']), vuln['cvss3_base_score'] + vuln['cvss3_vector'] + 
            ('Description', vuln['description']),   
            ('CVSS:',  cvss3 +"\n"+ cvss2 ),
            ('Affected IPs or URLs/Ports:', affected),
            ('Severity:', vuln['risk_factor']),
            ('Status:', 'Open'),
            ('Reference',see_also),
            ('Recommended actions:',vuln['solution']),
            ('Evidence:',plugin_output)
        )
### In this line is table created
        table = self.document.add_table(rows=1, cols=2)
##Styles REquired by table.
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'NAME'
        hdr_cells[1].text = vuln['pluginName']
        
        for det,desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = det 
            row_cells[1].text = desc
        
        # Adding style to the table
        table.style = 'Colorful List'
        self.document.add_page_break()


if __name__ == "__main__":
    # Banner print and arguments validation
    print(banner)
    parser = argparse.ArgumentParser(description='This script will load nessus v2 files and generate a output.docx with the data')
    parser.add_argument('-i',dest="ipath",required=True, help="-i <PATHFILES> path for files to parser")
    parser.add_argument('-o',dest="ofile",required=True ,help="Name file for result merge {fileout}.docx")
    parser.add_argument('-s',dest="severity",required=False, type = int, default = 1 ,help="Severity level corresponds to the numbers 0 through 4 ")

    args = parser.parse_args()
    if not args.ipath:
        print("[!] Specify a filename.docx and file output")
        sys.exit()
    if not args.ofile:
        print("[!] Specify the name outfile with -o")
        sys.exit()
    if args.severity < 0 or args.severity > 4:
        print("[!] Severity level corresponds to the numbers 0 through 4 ")
        sys.exit()

    nParser = NessusParser(args.severity)
    try:
        if os.path.isdir(args.ipath):
            if args.ipath.endswith('\\'):
                nessus_list = glob.glob(args.ipath + '*.nessus')
            else:
                nessus_list = glob.glob(args.ipath + '\*.nessus')
            print("Total files to process: ",len(nessus_list))
            if len(nessus_list) > 0:
                vulners = nParser.get_vulners_from_xmls(nessus_list)
                print("Total vulnerabilities found to be included in the report: ",len(vulners))
                report = Reporter(args.ofile)
                report.vulnreader(vulners)
                print ('[+] Saving results to: ' + args.ofile + '\n')
                print("[+] Done!")
            else:
                print ('\nError: No Nessus files were found in the supplied directory: "{0}"\n'.format(args.ipath))
                sys.exit(-1)
        else:
            print ('Error: "{0}" Is not a valid directory'.format(args.ipath))
            sys.exit(-1)
    except Exception as exc1:
        print("Error accesing the file: ",exc1)
